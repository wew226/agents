import os
import re
from pathlib import Path
from typing import Dict

import resend
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from agents import Agent, function_tool


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9_-]+", "_", value.strip())
    return cleaned.strip("_") or "learning_roadmap"


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(60, 60, 60)

    styles["Title"].font.name = "Georgia"
    styles["Title"].font.size = Pt(26)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = RGBColor(25, 38, 74)

    styles["Heading 1"].font.name = "Georgia"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)

    styles["Heading 2"].font.name = "Georgia"
    styles["Heading 2"].font.size = Pt(15)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(46, 84, 150)

    styles["Heading 3"].font.name = "Aptos"
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.bold = True
    styles["Heading 3"].font.color.rgb = RGBColor(90, 90, 90)


def _add_cover(doc: Document, goal: str, delivery_target: str, export_format: str) -> None:
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(goal).bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Personalized Learning Roadmap")
    run.font.name = "Aptos"
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(90, 90, 90)

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    labels = [("Delivery target", delivery_target), ("Export format", export_format)]
    for i, (label, value) in enumerate(labels):
        left = table.cell(i, 0)
        right = table.cell(i, 1)
        left.text = label
        right.text = value
        _set_cell_shading(left, "D9EAF7")
        for paragraph in left.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    doc.add_paragraph("")


def _render_markdown_to_docx(doc: Document, markdown_content: str) -> None:
    for raw_line in markdown_content.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("# "):
            doc.add_paragraph(stripped[2:].strip(), style="Heading 1")
            continue
        if stripped.startswith("## "):
            doc.add_paragraph(stripped[3:].strip(), style="Heading 2")
            continue
        if stripped.startswith("### "):
            doc.add_paragraph(stripped[4:].strip(), style="Heading 3")
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(stripped[2:].strip())
            continue
        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            p.add_run(re.sub(r"^\d+\.\s+", "", stripped))
            continue

        p = doc.add_paragraph()
        p.add_run(stripped)


def _beautify_docx(goal: str, markdown_content: str, delivery_target: str, export_format: str, path: Path) -> None:
    doc = Document()
    _style_document(doc)
    _add_cover(doc, goal, delivery_target, export_format)
    _render_markdown_to_docx(doc, markdown_content)

    section = doc.sections[-1]
    section.start_type = WD_SECTION.CONTINUOUS
    doc.save(path)


def _markdown_to_email_html(goal: str, markdown_content: str, file_path: str, delivery_target: str, export_format: str) -> str:
    html_lines: list[str] = []
    in_list = False

    def close_list() -> None:
        nonlocal in_list
        if in_list:
            html_lines.append("</ul>")
            in_list = False

    for raw_line in markdown_content.splitlines():
        line = raw_line.strip()
        if not line:
            close_list()
            continue

        if line.startswith("# "):
            close_list()
            html_lines.append(f"<h1>{line[2:].strip()}</h1>")
        elif line.startswith("## "):
            close_list()
            html_lines.append(f"<h2>{line[3:].strip()}</h2>")
        elif line.startswith("### "):
            close_list()
            html_lines.append(f"<h3>{line[4:].strip()}</h3>")
        elif line.startswith("- "):
            if not in_list:
                html_lines.append("<ul>")
                in_list = True
            html_lines.append(f"<li>{line[2:].strip()}</li>")
        else:
            close_list()
            html_lines.append(f"<p>{line}</p>")

    close_list()

    body = "\n".join(html_lines)
    return f"""<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <title>{goal}</title>
  <style>
    body {{
      font-family: Arial, sans-serif;
      background: #f5f7fb;
      color: #1f2937;
      margin: 0;
      padding: 32px 16px;
    }}
    .container {{
      max-width: 820px;
      margin: 0 auto;
      background: #ffffff;
      border-radius: 14px;
      padding: 36px 40px;
      box-shadow: 0 8px 30px rgba(15, 23, 42, 0.08);
    }}
    .eyebrow {{
      color: #2563eb;
      font-size: 12px;
      font-weight: 700;
      letter-spacing: 0.12em;
      text-transform: uppercase;
      margin-bottom: 10px;
    }}
    h1, h2, h3 {{
      color: #14213d;
      font-family: Georgia, serif;
    }}
    h1 {{
      font-size: 32px;
      margin-bottom: 18px;
      border-bottom: 1px solid #dbe4f0;
      padding-bottom: 12px;
    }}
    h2 {{
      margin-top: 32px;
      font-size: 22px;
      color: #1d4ed8;
    }}
    h3 {{
      margin-top: 22px;
      font-size: 16px;
      color: #475569;
      text-transform: uppercase;
      letter-spacing: 0.04em;
    }}
    p, li {{
      line-height: 1.8;
      font-size: 15px;
      color: #334155;
    }}
    .meta {{
      background: #f8fbff;
      border-left: 4px solid #60a5fa;
      padding: 16px 18px;
      margin: 18px 0 28px;
      border-radius: 8px;
    }}
    .footer {{
      margin-top: 32px;
      padding-top: 20px;
      border-top: 1px solid #e5e7eb;
      font-size: 14px;
      color: #64748b;
    }}
    code {{
      background: #eef2ff;
      padding: 2px 6px;
      border-radius: 4px;
      font-size: 13px;
    }}
  </style>
</head>
<body>
  <div class="container">
    <div class="eyebrow">Learning Roadmap Ready</div>
    <div class="meta">
      <strong>Delivery target:</strong> {delivery_target}<br>
      <strong>Export format:</strong> {export_format}<br>
      <strong>File path:</strong> <code>{file_path}</code>
    </div>
    {body}
    <div class="footer">
      Your complete roadmap has been generated and exported successfully.
    </div>
  </div>
</body>
</html>
"""


@function_tool
def export_document(
    goal: str,
    markdown_content: str,
    output_dir: str,
    export_format: str,
    delivery_target: str,
) -> Dict[str, str]:
    """Export the roadmap to DOCX, HTML, Markdown, or CSV and return file metadata."""
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    base = _safe_filename(goal)
    export_format = export_format.upper().strip()
    delivery_target = delivery_target.strip()

    if export_format == "DOCX":
        path = out_dir / f"{base}.docx"
        _beautify_docx(goal, markdown_content, delivery_target, export_format, path)
    elif export_format == "HTML":
        path = out_dir / f"{base}.html"
        html = _markdown_to_email_html(goal, markdown_content, "", delivery_target, export_format)
        path.write_text(html, encoding="utf-8")
    elif export_format == "MARKDOWN":
        path = out_dir / f"{base}.md"
        path.write_text(markdown_content, encoding="utf-8")
    elif export_format == "CSV":
        path = out_dir / f"{base}.csv"
        rows = ["section,content"]
        for line in markdown_content.splitlines():
            safe = line.replace('"', '""')
            rows.append(f"roadmap,\"{safe}\"")
        path.write_text("\n".join(rows), encoding="utf-8")
    else:
        raise ValueError(f"Unsupported export format: {export_format}")

    return {
        "status": "success",
        "file_path": str(path.resolve()),
        "delivery_target": delivery_target,
        "export_format": export_format,
    }


@function_tool
def send_email(subject: str, html_body: str) -> Dict[str, str]:
    """Send an email with the given subject and HTML body."""
    resend.api_key = os.environ["RESEND_API_KEY"]
    resend.Emails.send({
        "from": "onboarding@resend.dev",
        "to": os.environ["RESEND_TO_EMAIL"],
        "subject": subject,
        "html": html_body,
    })
    return {"status": "success"}


INSTRUCTIONS = """
You deliver a learning roadmap after it has been written.

You will receive:
- the learning goal
- the delivery target
- the export format
- the output directory
- the roadmap markdown

Your job:
1. Export the roadmap using the export_document tool.
2. Always send an email after export is complete.
3. The email must contain the full generated roadmap content, not just a short confirmation.
4. The email must look polished and include:
   - the learning goal
   - the delivery target
   - the export format
   - the exported file path
   - the full roadmap rendered as readable HTML
5. Use a subject line exactly like:
   "Your learning roadmap is ready: <goal>"
6. Return a concise delivery summary with file path and email status.
"""

delivery_agent = Agent(
    name="DeliveryAgent",
    instructions=INSTRUCTIONS,
    tools=[export_document, send_email],
    model="gpt-4o-mini",
)
